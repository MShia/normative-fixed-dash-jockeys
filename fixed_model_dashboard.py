"""
Neuropsychological Normative Calculator
Requirements: pip install streamlit pandas numpy torch scipy python-docx plotly - education years + save results logic
"""

import streamlit as st
import pandas as pd
import numpy as np
import torch
import torch.nn as nn
import pickle
import os
from scipy.stats import norm
from pathlib import Path
import json
from typing import Dict, List, Any
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import plotly.graph_objects as go

# Set page config
st.set_page_config(
    page_title="Neuropsychological Normative Calculator",
    page_icon="🏇",  # Horse racing emoji as favicon
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional appearance
st.markdown("""
<style>
    /* Hide default Streamlit padding */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 1rem;
        max-width: 100%;
    }
    
    /* Main header styling */
    .main-header {
        font-size: 2rem;
        font-weight: 600;
        color: #1a1a1a;
        text-align: left;
        margin-bottom: 1.5rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e0e0e0;
    }
    
    /* Section headers */
    .section-header {
        font-size: 1.1rem;
        font-weight: 600;
        color: #2c3e50;
        margin-top: 1rem;
        margin-bottom: 0.8rem;
    }
    
    /* Result cards with modern design */
    .result-card-main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 0.5rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .result-card-flagged {
        background: linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 0.5rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .result-card-safe {
        background: linear-gradient(135deg, #00d2d3 0%, #229954 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 0.5rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    /* Result card container for visual display */
    .result-visual-card {
        background: white;
        border-radius: 12px;
        padding: 1rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        text-align: center;
    }
    
    /* Method cards */
    .method-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 3px solid #3498db;
        margin: 0.5rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    .method-card-flagged {
        background-color: #fff5f5;
        padding: 1rem;
        border-radius: 8px;
        border-left: 3px solid #e74c3c;
        margin: 0.5rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    .method-card-safe {
        background-color: #f0fff4;
        padding: 1rem;
        border-radius: 8px;
        border-left: 3px solid #27ae60;
        margin: 0.5rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    /* Info boxes */
    .info-box {
        background-color: #f0f4f8;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
    
    .info-box-alert {
        background-color: #fff5f5;
        border: 1px solid #feb2b2;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .info-box-success {
        background-color: #f0fff4;
        border: 1px solid #9ae6b4;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    /* Compact form styling */
    .stNumberInput > div > div > input {
        font-size: 0.9rem;
    }
    
    .stSelectbox > div > div > div {
        font-size: 0.9rem;
    }
    
    /* Button styling */
    .stButton > button {
        height: 2.2rem;
        font-size: 0.9rem;
        font-weight: 500;
        border-radius: 6px;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    /* Reduce spacing between elements */
    .element-container {
        margin-bottom: 0.5rem;
    }
    
    /* Metric styling */
    .metric-container {
        background: white;
        border-radius: 8px;
        padding: 0.8rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        text-align: center;
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        margin: 0;
    }
    
    .metric-label {
        font-size: 0.85rem;
        color: #666;
        margin: 0;
    }
    
    /* Tooltip styling */
    .tooltip-container {
        position: relative;
        display: inline-block;
        cursor: help;
    }
    
    .tooltip-text {
        visibility: hidden;
        width: 250px;
        background-color: #333;
        color: #fff;
        text-align: left;
        border-radius: 6px;
        padding: 8px 10px;
        position: absolute;
        z-index: 1;
        bottom: 125%;
        left: 0;
        margin-left: -10px;
        opacity: 0;
        transition: opacity 0.3s;
        font-size: 0.85rem;
        line-height: 1.4;
    }
    
    .tooltip-text::after {
        content: "";
        position: absolute;
        top: 100%;
        left: 20px;
        margin-left: -5px;
        border-width: 5px;
        border-style: solid;
        border-color: #333 transparent transparent transparent;
    }
    
    .tooltip-container:hover .tooltip-text {
        visibility: visible;
        opacity: 1;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Responsive design */
    @media (max-width: 768px) {
        .main-header {
            font-size: 1.5rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# Neural Network Architecture
class MultiQuantileRegressionNet(nn.Module):
    def __init__(self, input_dim, num_quantiles=99):
        super().__init__()
        self.model = nn.Sequential(
            nn.Linear(input_dim, 64),
            nn.ReLU(),
            nn.BatchNorm1d(64),
            nn.Linear(64, 64),
            nn.ReLU(),
            nn.Linear(64, 64),
            nn.ReLU(),
            nn.Dropout(0.02),
            nn.Linear(64, num_quantiles)
        )

    def forward(self, x):
        return self.model(x)

# Model Manager Class
class ModelManager:
    def __init__(self, models_dir="models"):
        self.models_dir = Path(models_dir)
        self.quantiles = np.linspace(0.01, 0.99, 99)
        self.available_scores = []
        self.models = {}
        self.metadata = {}
        self.coverage_data = {}
        self.load_models()
    
    def load_models(self):
        """Load all pre-trained models and metadata"""
        try:
            # Load metadata
            metadata_path = self.models_dir / "model_metadata.json"
            if metadata_path.exists():
                with open(metadata_path, 'r') as f:
                    self.metadata = json.load(f)
                    self.available_scores = list(self.metadata.keys())
            
            # Load coverage data if available
            coverage_path = self.models_dir / "coverage_data.json"
            if coverage_path.exists():
                with open(coverage_path, 'r') as f:
                    self.coverage_data = json.load(f)
            
            # Load models for each score
            for score in self.available_scores:
                self.models[score] = {}
                
                # Load LR model
                lr_path = self.models_dir / f"{score}_lr_model.pkl"
                if lr_path.exists():
                    with open(lr_path, 'rb') as f:
                        self.models[score]['LR'] = pickle.load(f)
                
                # Load LQR models (99 models)
                lqr_models = []
                for i, q in enumerate(self.quantiles):
                    lqr_path = self.models_dir / f"{score}_lqr_q{i:02d}.pkl"
                    if lqr_path.exists():
                        with open(lqr_path, 'rb') as f:
                            lqr_models.append(pickle.load(f))
                if lqr_models:
                    self.models[score]['LQR'] = lqr_models
                
                # Load NNQR model and scaler
                nnqr_path = self.models_dir / f"{score}_nnqr_model.pth"
                scaler_path = self.models_dir / f"{score}_scaler.pkl"
                
                if nnqr_path.exists() and scaler_path.exists():
                    input_dim = self.metadata[score].get('input_dim', 3)
                    model = MultiQuantileRegressionNet(input_dim)
                    model.load_state_dict(torch.load(nnqr_path, map_location='cpu'))
                    model.eval()
                    
                    with open(scaler_path, 'rb') as f:
                        scaler = pickle.load(f)
                    
                    self.models[score]['NNQR'] = {
                        'model': model,
                        'scaler': scaler
                    }
                    
        except Exception as e:
            st.error(f"Error loading models: {str(e)}")
    
    def predict_lr(self, score, features, actual_score):
        """Linear Regression prediction"""
        if score not in self.models or 'LR' not in self.models[score]:
            return None
        
        try:
            model = self.models[score]['LR']
            predicted_score = model['intercept'] + np.dot(features, model['coefficients'])
            z_score = (actual_score - predicted_score) / model['std']
            percentile = norm.cdf(z_score) * 100
            # Round to avoid floating point precision issues
            percentile = np.round(percentile, 2)
            return np.clip(percentile, 1, 99)
        except Exception as e:
            st.error(f"LR prediction error: {str(e)}")
            return None
    
    def predict_lqr(self, score, features, actual_score):
        """Linear Quantile Regression prediction"""
        if score not in self.models or 'LQR' not in self.models[score]:
            return None
        
        try:
            models = self.models[score]['LQR']
            predictions = []
            for model in models:
                pred = model['intercept'] + np.dot(features, model['coefficients'])
                predictions.append(pred)
            
            predictions = np.array(predictions)
            greater_idx = np.where(predictions > actual_score)[0]
            if len(greater_idx) > 0:
                percentile = self.quantiles[greater_idx[0]] * 100
            else:
                percentile = 99
            
            # Round to avoid floating point precision issues
            percentile = np.round(percentile, 2)
            return np.clip(percentile, 1, 99)
        except Exception as e:
            st.error(f"LQR prediction error: {str(e)}")
            return None
    
    def predict_nnqr(self, score, features, actual_score):
        """Neural Network Quantile Regression prediction"""
        if score not in self.models or 'NNQR' not in self.models[score]:
            return None
        
        try:
            model = self.models[score]['NNQR']['model']
            scaler = self.models[score]['NNQR']['scaler']
            
            features_scaled = scaler.transform(features.reshape(1, -1))
            features_tensor = torch.FloatTensor(features_scaled)
            
            with torch.no_grad():
                predictions = model(features_tensor).numpy().flatten()
            
            greater_idx = np.where(predictions > actual_score)[0]
            if len(greater_idx) > 0:
                percentile = self.quantiles[greater_idx[0]] * 100
            else:
                percentile = 99
            
            # Round to avoid floating point precision issues
            percentile = np.round(percentile, 2)
            return np.clip(percentile, 1, 99)
        except Exception as e:
            st.error(f"NNQR prediction error: {str(e)}")
            return None
    
    def get_coverage_offset(self, score, method, percentile):
        """Get coverage offset for a specific method and percentile"""
        if score not in self.coverage_data or method not in self.coverage_data[score]:
            return 0.01
        
        try:
            coverage = self.coverage_data[score][method]
            quantile_idx = min(98, max(0, int(percentile) - 1))
            theoretical_quantile = self.quantiles[quantile_idx]
            empirical_coverage = coverage[quantile_idx]
            offset = abs(empirical_coverage - theoretical_quantile)
            return offset
        except:
            return 0.01
    
    def calculate_percentile_with_best_method(self, score, actual_score, features, min_percentile_constraint=None, use_max_percentile=False):
        """Calculate percentile using all methods and select best based on coverage offset or max percentile
        
        Args:
            score: Test type
            actual_score: The actual test score
            features: Demographic features
            min_percentile_constraint: If provided, only select methods that give percentile >= this value
            use_max_percentile: If True, select method with highest percentile instead of best coverage
        """
        results = {}
        offsets = {}
        
        # Get predictions from all methods
        lr_percentile = self.predict_lr(score, features, actual_score)
        if lr_percentile is not None:
            results['LR'] = lr_percentile
            offsets['LR'] = self.get_coverage_offset(score, 'LR', lr_percentile)
        
        lqr_percentile = self.predict_lqr(score, features, actual_score)
        if lqr_percentile is not None:
            results['LQR'] = lqr_percentile
            offsets['LQR'] = self.get_coverage_offset(score, 'LQR', lqr_percentile)
        
        nnqr_percentile = self.predict_nnqr(score, features, actual_score)
        if nnqr_percentile is not None:
            results['NNQR'] = nnqr_percentile
            offsets['NNQR'] = self.get_coverage_offset(score, 'NNQR', nnqr_percentile)
        
        # Apply minimum percentile constraint if provided
        if min_percentile_constraint is not None:
            # Filter out methods that would create non-monotonic results
            valid_results = {method: pct for method, pct in results.items() 
                           if pct >= min_percentile_constraint}
            valid_offsets = {method: offset for method, offset in offsets.items() 
                           if method in valid_results}
            
            # If no methods satisfy the constraint, fall back to highest percentile
            if not valid_results:
                # Select method with highest percentile to maintain monotonicity
                best_method = max(results, key=results.get)
                best_percentile = results[best_method]
                best_offset = offsets[best_method]
                return results, offsets, best_method, best_percentile, best_offset
            
            # Use valid results/offsets for selection
            results_for_selection = valid_results
            offsets_for_selection = valid_offsets
        else:
            results_for_selection = results
            offsets_for_selection = offsets
        
        # Find best method based on selection criteria
        if results_for_selection:
            if use_max_percentile:
                # For maximum score, select method with highest percentile
                best_method = max(results_for_selection, key=results_for_selection.get)
            else:
                # Normal case: select method with minimum offset
                best_method = min(offsets_for_selection, key=offsets_for_selection.get)
            
            best_percentile = results[best_method]
            best_offset = offsets[best_method]
        else:
            best_method = None
            best_percentile = None
            best_offset = None
        
        return results, offsets, best_method, best_percentile, best_offset
    
    def evaluate_flags(self, results, threshold, best_method=None, best_percentile=None):
        """Evaluate flagging for each method and agreement confidence"""
        flags = {}
        for method, percentile in results.items():
            # Use < threshold consistently for all methods
            # This means a score AT the threshold (e.g., 5th percentile when threshold is 5) is NOT flagged
            flags[method] = percentile < threshold
        
        flag_count = sum(flags.values())
        total_methods = len(flags)
        
        # If best_method and best_percentile provided, use new logic
        if best_method is not None and best_percentile is not None:
            best_flagged = best_percentile < threshold  # Consistent with above
            
            # Determine majority opinion
            majority_says_flag = flag_count > total_methods / 2
            
            # Check if best method agrees with majority
            best_agrees_with_majority = (best_flagged == majority_says_flag)
            
            if flag_count == 0:
                agreement_level = "Unanimous Pass"
                agreement_confidence = "High Agreement"
                agreement_color = "#27ae60"  # Green - all agree no flag
                agreement_ratio = f"0/{total_methods} flagged"
            elif flag_count == total_methods:
                agreement_level = "Unanimous Flag"
                agreement_confidence = "High Agreement"
                agreement_color = "#27ae60"  # Green - all agree flag
                agreement_ratio = f"{flag_count}/{total_methods} flagged"
            else:
                # Mixed results - check if best agrees with majority
                if best_agrees_with_majority:
                    agreement_confidence = "Moderate Agreement"
                    agreement_color = "#27ae60"  # Green - best agrees with majority
                else:
                    agreement_confidence = "Low Agreement"
                    agreement_color = "#e74c3c"  # Red - best disagrees with majority
                
                if majority_says_flag:
                    agreement_level = "Majority Flag"
                else:
                    agreement_level = "Majority Pass"
                agreement_ratio = f"{flag_count}/{total_methods} flagged"
        else:
            # Fallback to old logic for backward compatibility
            agreement_confidence = "Not Available"  # Default when best method not provided
            
            if flag_count == 0:
                agreement_level = "No Flag"
                agreement_color = "#27ae60"  # Green
                agreement_ratio = f"0/{total_methods}"
            elif flag_count == total_methods:
                agreement_level = "All Flag"
                agreement_color = "#e74c3c"  # Red
                agreement_ratio = f"{flag_count}/{total_methods}"
            else:
                if flag_count > total_methods / 2:
                    agreement_level = "Majority Flag"
                    agreement_color = "#e74c3c"  # Red
                else:
                    agreement_level = "Majority No Flag"
                    agreement_color = "#27ae60"  # Green
                agreement_ratio = f"{flag_count}/{total_methods}"
        
        # Always return 5 values
        return flags, agreement_level, agreement_color, agreement_ratio, agreement_confidence

# Initialize model manager
@st.cache_resource
def load_model_manager():
    return ModelManager()

# Initialize session state for storing all test results
def init_session_state():
    if 'all_test_results' not in st.session_state:
        st.session_state.all_test_results = []
    if 'test_counter' not in st.session_state:
        st.session_state.test_counter = 0

# Function to create percentile ring visualization
def create_percentile_ring(percentile, flag_status,h,w,f):
    """Create a circular ring chart showing percentile visually"""
    # Calculate how many segments to fill
    segments = 20  # Total number of segments for finer granularity
    filled_segments = int((percentile / 100) * segments)
    remaining_segments = segments - filled_segments
    
    # Create values for the pie chart
    values = [1] * segments
    
    # Create gradient colors based on percentile position
    colors = []
    for i in range(segments):
        segment_percentile = (i + 1) * (100 / segments)
        if segment_percentile <= percentile:
            # Color the filled portion based on flag_status
            if flag_status == True:
                colors.append('#e74c3c')  # Red for flagged
                
            elif flag_status == False:
                colors.append('#27ae60')  # Green for pass
                
            else:
                colors.append("#d3d3d3")  # Gray for unknown status
        else:
            # Empty segments in light gray
            colors.append("#d3d3d3")  # Light gray for unfilled portion
    
    # Create the pie chart
    fig = go.Figure(go.Pie(
        values=values,
        marker=dict(colors=colors, line=dict(color='white', width=2)),
        hole=0.55,
        direction='clockwise',
        sort=False,
        textinfo='none',
        hoverinfo='none',
        showlegend=False
    ))
    
    # Update layout
    fig.update_layout(
        showlegend=False,
        margin=dict(t=0, b=0, l=0, r=0),
        width=w,
        height=h,
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)'
    )
    
    # Determine text color based on percentile
    if flag_status == True:
        text_color = '#e74c3c'
    elif flag_status == False:
        text_color = '#27ae60'
    else:
        text_color = "#070707"
    
    # Add percentile text in the center
    fig.add_annotation(
        text=f"<b>{percentile:.0f}</b><br><span style='font-size:{f}px'>%</span>",
        x=0.5, y=0.5,
        xref="paper", yref="paper",
        showarrow=False,
        font=dict(size=f, color=text_color, family="Arial")
    )
    
    return fig

# Function to generate Word report
def generate_word_report(all_results, demographics):
    """Generate a comprehensive Word report for all test results"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Neuropsychological Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report metadata
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Total Tests Evaluated: {len(all_results)}")
    doc.add_paragraph()
    
    # Subject Information section
    doc.add_heading('Subject Information', level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light List Accent 1'
    
    # Populate subject info
    info_data = [
        ('Age', f"{demographics['age']} years"),
        ('Sex', demographics['sex']),
        ('Education', f"{demographics['education']} years"),
        ('Assessment Date', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    
    # Test Results Summary
    doc.add_heading('Test Results Summary', level=1)
    
    # Create summary table
    summary_table = doc.add_table(rows=len(all_results)+1, cols=5)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Test', 'Raw Score', 'Percentile',  'Flag Status', 'Method Agreement']
    for i, header in enumerate(headers):
        cell = summary_table.cell(0, i)
        cell.text = header
        # Bold headers
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
    
    # Populate summary table
    for i, result in enumerate(all_results, 1):
        summary_table.cell(i, 0).text = result['test_name'].replace('_raw', '').upper()
        summary_table.cell(i, 1).text = f"{result['actual_score']:.2f}"
        summary_table.cell(i, 2).text = f"{result['best_percentile']:.0f}th"
        # summary_table.cell(i, 3).text = result['interpretation']
        
        # Flag status with color
        flag_cell = summary_table.cell(i, 3)
        flag_cell.text = result['flag_status']
        if result['flag_status'] == 'FLAGGED':
            run = flag_cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(231, 76, 60)  # Red
        else:
            run = flag_cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(39, 174, 96)  # Green
        
        summary_table.cell(i, 4).text = result['agreement_ratio']
    
    doc.add_page_break()
    
    # Detailed Results for each test
    doc.add_heading('Detailed Test Results', level=1)
    
    for idx, result in enumerate(all_results, 1):
        # Test header
        doc.add_heading(f"{idx}. {result['test_name'].replace('_raw', '').upper()}", level=2)
        
        # Create detailed table for this test
        detail_table = doc.add_table(rows=5, cols=2)
        detail_table.style = 'Light List Accent 1'
        
        details = [
            ('Raw Score', f"{result['actual_score']:.2f}"),
            ('Best Percentile (Method)', f"{result['best_percentile']:.0f}th ({result['best_method']})"),
            # ('Clinical Interpretation', result['interpretation']),
            ('Flag Status (Threshold)', f"(< {result['threshold']}th percentile)"),
            ('Method Agreement', f"{result['agreement_level']} ({result['agreement_ratio']})")
        ]
        
        for i, (label, value) in enumerate(details):
            detail_table.cell(i, 0).text = label
            detail_table.cell(i, 1).text = value
        
        # Method-specific results
        doc.add_paragraph()
        doc.add_paragraph('Method-Specific Results:', style='Heading 3')
        
        method_table = doc.add_table(rows=len(result['all_methods'])+1, cols=4)
        method_table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['Method', 'Percentile', 'Flag Status', 'Coverage Offset']
        for i, header in enumerate(headers):
            cell = method_table.cell(0, i)
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        
        # Method results
        for i, (method, data) in enumerate(result['all_methods'].items(), 1):
            method_table.cell(i, 0).text = method
            method_table.cell(i, 1).text = f"{data['percentile']:.1f}th"
            method_table.cell(i, 2).text = data['flag']
            method_table.cell(i, 3).text = f"{data['offset']:.4f}"
        
        if idx < len(all_results):
            doc.add_paragraph()
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
    
    # Clinical Summary
    doc.add_page_break()
    doc.add_heading('Clinical Summary', level=1)
    
    # Count flagged tests
    flagged_tests = [r for r in all_results if r['flag_status'] == 'FLAGGED']
    
    summary_text = f"""
    This neuropsychological assessment included {len(all_results)} cognitive tests. 
    Of these, {len(flagged_tests)} test(s) fell below the clinical threshold.
    
    """
    
    if flagged_tests:
        summary_text += "The following tests showed clinically significant findings:\n"
        for test in flagged_tests:
            summary_text += f"• {test['test_name'].replace('_raw', '').upper()}: {test['best_percentile']:.0f}th percentile\n"
    else:
        summary_text += "All test scores fell within or above the expected range for the subject's demographic profile."
    
    doc.add_paragraph(summary_text)
    
    # Disclaimer
    doc.add_paragraph()
    doc.add_heading('Disclaimer', level=2)
    disclaimer = doc.add_paragraph(
        "This report is generated by an automated normative calculator and should be interpreted "
        "by qualified professionals in conjunction with clinical observations and other relevant information. "
        "The percentiles are calculated based on demographic-adjusted normative data."
    )
    disclaimer.style = 'Intense Quote'
    
    # Convert to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

# Main Application
def main():
    # Initialize session state
    init_session_state()
    
    # Header with custom logo
    with open("dashlogo1.svg", "r") as f:
        logo_svg = f.read()
    
    col1, col2 = st.columns([1, 8])
    with col1:
        st.markdown(f'<div style="padding-top: 0.5rem;">{logo_svg}</div>', unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="main-header">Comparative Neuropsychological Percentile Estimation Based on Irish Jockeys Data</div>', unsafe_allow_html=True)
    
    # Research description
    st.markdown("""
    <div style="background-color: #e8f4f8; border-left: 4px solid #1f77b4; padding: 1rem; 
                border-radius: 8px; margin-bottom: 1.5rem; font-size: 0.95rem;">
        <strong>📚 About this Research Tool:</strong><br>
        This dashboard resulted from research funding by HRB Ireland’s Secondary Data Analysis Grant, presents regression-based normative models 
        for neuropsychological tests in Irish jockeys (2010–2024) to support concussion assessment. Percentile estimates 
        are generated using age, sex, and education via three methods: Linear Regression (LR), 
        Linear Quantile Regression (LQR), and Neural Network Quantile Regression (NNQR). The best estimate 
        is selected based on empirical coverage offset, and agreement confidence reflects consistency across the three methods.
    </div>
    """, unsafe_allow_html=True)
    
    # Load models
    model_manager = load_model_manager()
    
    if not model_manager.available_scores:
        st.error("❌ No pre-trained models found. Please ensure models are properly installed.")
        return
    
    # Create layout with sidebar and main content
    with st.sidebar:
        st.markdown("### 👤 Subject Demographics")
        
        col1, col2 = st.columns(2)
        with col1:
            age = st.number_input("Age", min_value=16, max_value=56, value=35, step=1)
        with col2:
            sex = st.selectbox("Sex", ["Male", "Female"])
        
        # Calculate maximum realistic education years based on age
        # Assuming education starts at age 6, max education = age - 6
        max_education = min(22, age - 6)  # Cap at 22 years maximum
        
        # Show warning if current education value would be invalid
        if 'prev_age' not in st.session_state:
            st.session_state.prev_age = age
        
        education = st.number_input(
            "Education (years)", 
            min_value=8, 
            max_value=max_education, 
            value=min(12, max_education),  # Default to 12 or max allowed
            step=1,
            help=f"Maximum education years for age {age} is {max_education} (assuming education starts at age 6)"
        )
        
        # Show validation message if needed
        if education > max_education:
            st.error(f"⚠️ Education years cannot exceed {max_education} for age {age}")
        
        sex_numeric = 1 if sex == "Male" else 0
        features = np.array([age, sex_numeric, education])
        
        # Store demographics in session state
        st.session_state.demographics = {
            'age': age,
            'sex': sex,
            'education': education
        }
        
        st.markdown(f"""
        <div class="info-box">
        <strong>Current Profile</strong><br>
        Age: {age} | Sex: {sex} | Education: {education}y
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 📊 Available Tests")
        
        # Test descriptions with tooltips
        test_descriptions = {
            "DSF": "A test of short-term memory where individuals repeat a sequence of numbers in the same order as presented.",
            "DSB": "A test of working memory that requires individuals to repeat a sequence of numbers in reverse order.",
            "SOC": "A test measuring how quickly and accurately a person can understand and respond to simple written information.",
            "SDMT": "A test assessing attention, visual scanning, and processing speed, where individuals match symbols to numbers using a reference key."
        }
        
        # Display tests with tooltips
        st.markdown(f"""
        <ul style="list-style-type: none; padding-left: 0;">
            <li style="margin-bottom: 0.5rem;">
                <div class="tooltip-container">
                    <strong>DSF</strong>: Digit Span Forward
                    <span class="tooltip-text">{test_descriptions['DSF']}</span>
                </div>
            </li>
            <li style="margin-bottom: 0.5rem;">
                <div class="tooltip-container">
                    <strong>DSB</strong>: Digit Span Backward
                    <span class="tooltip-text">{test_descriptions['DSB']}</span>
                </div>
            </li>
            <li style="margin-bottom: 0.5rem;">
                <div class="tooltip-container">
                    <strong>SOC</strong>: Speed of Comprehension
                    <span class="tooltip-text">{test_descriptions['SOC']}</span>
                </div>
            </li>
            <li style="margin-bottom: 0.5rem;">
                <div class="tooltip-container">
                    <strong>SDMT</strong>: Symbol Digit Modalities
                    <span class="tooltip-text">{test_descriptions['SDMT']}</span>
                </div>
            </li>
        </ul>
        """, unsafe_allow_html=True)
        
        st.markdown(f"<small>✅ {len(model_manager.available_scores)} models loaded</small>", unsafe_allow_html=True)
        
        # Session summary
        if st.session_state.all_test_results:
            st.markdown("---")
            st.markdown("### 📋 Session Summary")
            st.markdown(f"**Tests Completed:** {len(st.session_state.all_test_results)}")
            
            for i, result in enumerate(st.session_state.all_test_results):
                test_name = result['test_name'].replace('_raw', '').upper()
                flag_emoji = "🔴" if result['flag_status'] == 'FLAGGED' else "🟢"
                st.markdown(f"{flag_emoji} {test_name}: {result['best_percentile']:.0f}th")
            
            # Report generation section
            st.markdown("---")
            st.markdown("### 📄 Generate Report")
            
            # Generate report buffer
            report_buffer = generate_word_report(
                st.session_state.all_test_results,
                st.session_state.demographics
            )
            
            # Download button (always visible when there are results)
            st.download_button(
                label="📥 Download Full Report (Word)",
                data=report_buffer,
                file_name=f"neuropsych_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            if st.button("🗑️ Clear Session", use_container_width=True):
                st.session_state.all_test_results = []
                st.session_state.test_counter = 0
                st.rerun()
    
    # Main content area - single column layout for better space usage
    # Input section
    st.markdown('<div class="section-header">📝 Test Configuration</div>', unsafe_allow_html=True)
    
    # Define score ranges for each test
    score_ranges = {
        'dsf_raw': {'min': 5, 'max': 16},
        'DSF_raw': {'min': 5, 'max': 16},
        'dsb_raw': {'min': 3, 'max': 14},
        'DSB_raw': {'min': 3, 'max': 14},
        'soc_raw': {'min': 11, 'max': 100},
        'SOC_raw': {'min': 11, 'max': 100},
        'sdmt_raw': {'min': 15, 'max': 100},
        'SDMT_raw': {'min': 15, 'max': 100}
    }
    
    col1, col2, col3, col4 = st.columns([3, 2, 2, 2])
    
    with col1:
    # Define mapping for test names
        test_name_mapping = {
            'DSF_raw': 'Digit Span Forward (DSF)',
            'DSB_raw': 'Digit Span Backward (DSB)',
            'SOC_raw': 'Speed of Comprehension (SOC)',
            'SDMT_raw': 'Symbol Digit Modalities (SDMT)'
        }
    
        selected_score = st.selectbox(
            "Select Test",
            model_manager.available_scores,
            format_func=lambda x: test_name_mapping.get(x, x.replace('_raw', '').upper())
        )

    with col2:
        threshold = st.number_input(
            "Flag Threshold (%ile)",
            min_value=1,
            max_value=99,
            value=5,
            step=1,
            help="Flag scores below this percentile"
        )    
    
    with col3:
        # Get the range for the selected test
        min_score = score_ranges.get(selected_score, {}).get('min', 0)
        max_score = score_ranges.get(selected_score, {}).get('max', 100)
        
        actual_score = st.number_input(
            "Actual Score",
            min_value=float(min_score),
            max_value=float(max_score),
            value=float(min_score),
            step=1.0,
            help=f"Enter raw test score (Range: {min_score} - {max_score})"
        )
    

    
    with col4:
        calculate_btn = st.button("🚀 Estimate", type="primary", use_container_width=True)
    
    # Results section
    if calculate_btn and actual_score is not None:
        with st.spinner("Calculating..."):
            # Check if we have a previous calculation for the same test and demographics
            min_percentile_constraint = None
            
            if 'calculation_history' not in st.session_state:
                st.session_state.calculation_history = []
            
            # Look for previous calculations with same test and demographics but lower scores
            for prev_calc in st.session_state.calculation_history:
                if (prev_calc['test'] == selected_score and 
                    prev_calc['features'] == features.tolist() and
                    prev_calc['score'] < actual_score):
                    # If there's a previous calculation with lower score, use its percentile as constraint
                    if min_percentile_constraint is None or prev_calc['percentile'] > min_percentile_constraint:
                        min_percentile_constraint = prev_calc['percentile']
            
            # Calculate with constraint if applicable
            results, offsets, best_method, best_percentile, best_offset = model_manager.calculate_percentile_with_best_method(
                selected_score, actual_score, features, min_percentile_constraint=min_percentile_constraint
            )
            
            if results:
                flags, agreement_level, agreement_color, agreement_ratio, agreement_confidence = model_manager.evaluate_flags(
                    results, threshold, best_method=best_method, best_percentile=best_percentile
                )
                
                # Add current calculation to history
                st.session_state.calculation_history.append({
                    'test': selected_score,
                    'features': features.tolist(),
                    'score': actual_score,
                    'percentile': best_percentile,
                    'method': best_method
                })
                
                # Keep only last 50 calculations to avoid memory issues
                if len(st.session_state.calculation_history) > 50:
                    st.session_state.calculation_history = st.session_state.calculation_history[-50:]
                
                # Store current calculation in session state
                st.session_state.current_results = {
                    'results': results,
                    'offsets': offsets,
                    'best_method': best_method,
                    'best_percentile': best_percentile,
                    'best_offset': best_offset,
                    'flags': flags,
                    'agreement_level': agreement_level,
                    'agreement_color': agreement_color,
                    'agreement_ratio': agreement_ratio,
                    'agreement_confidence': agreement_confidence,
                    'selected_score': selected_score,
                    'actual_score': actual_score,
                    'threshold': threshold
                }
            else:
                st.error("❌ Unable to estimate percentiles. Please check your inputs.")
    
    # Display results if they exist in session state
    if 'current_results' in st.session_state and st.session_state.current_results:
        results = st.session_state.current_results['results']
        offsets = st.session_state.current_results['offsets']
        best_method = st.session_state.current_results['best_method']
        best_percentile = st.session_state.current_results['best_percentile']
        best_offset = st.session_state.current_results['best_offset']
        flags = st.session_state.current_results['flags']
        agreement_level = st.session_state.current_results['agreement_level']
        agreement_color = st.session_state.current_results['agreement_color']
        agreement_ratio = st.session_state.current_results['agreement_ratio']
        agreement_confidence = st.session_state.current_results.get('agreement_confidence', 'Moderate Agreement')
        current_score = st.session_state.current_results['selected_score']
        current_actual = st.session_state.current_results['actual_score']
        current_threshold = st.session_state.current_results['threshold']
        
        # Determine overall status
        best_flagged = best_percentile < current_threshold
        
        # Results header
        st.markdown('<div class="section-header">📊 Results Summary</div>', unsafe_allow_html=True)
        
        # Main result display
        col1, col2, col3 = st.columns([2, 3, 2])
        
        with col1:
            with st.container(border=True):
                # Best result header
                st.markdown("#### 🏆 Best Estimate")
                
                # Create percentile ring visualization
                ring_fig = create_percentile_ring(best_percentile, best_flagged, h=180, w=180, f=25)
                st.plotly_chart(ring_fig, use_container_width=True, config={'displayModeBar': False})
                
                # Method and status below the ring
                flag_text = "⚠️ FLAGGED" if best_flagged else "✅ PASS"
                flag_color = "#e74c3c" if best_flagged else "#27ae60"
                
                st.markdown(f"""
                <div style="text-align: center; margin-top: -20px;">
                    <div style="font-size: 1.2rem; font-weight: bold; color: #333; margin-bottom: 10px;">{best_method}</div>
                    <div style="font-size: 1.1rem; color: {flag_color}; font-weight: bold;">{flag_text}</div>
                </div>
        """, unsafe_allow_html=True)
        
        with col2:
            with st.container(border=True):
                # Method comparison header
                st.markdown("#### 📊 Method Comparison")
                
                # Create 3 columns for methods
                method_cols = st.columns(3)
                
                for i, (method, percentile) in enumerate(results.items()):
                    with method_cols[i]:
                        method_flagged = flags[method]

                        # Create the ring chart
                        ring_figure = create_percentile_ring(percentile, method_flagged, h=80, w=80, f=12)
                        
                        # Add unique key to avoid duplicate ID error
                        st.plotly_chart(
                            ring_figure, 
                            config={'displayModeBar': False},
                            key=f"ring_chart_{method}_{i}"  # Unique key for each chart
                        )
                        
                        # Method name and status
                        flag_text = "⚠️ FLAGGED" if method_flagged else "✅ PASS"
                        flag_color = "#e74c3c" if method_flagged else "#27ae60"
                        
                        st.markdown(f"""
                        <div style="text-align: center; margin-top: -20px;">
                            <div style="font-size: 1.2rem; text-align: center; font-weight: bold; color: #333; margin-bottom: 10px;">{method}</div>
                            <div style="font-size: 1.1rem; color: {flag_color}; font-weight: bold;">{flag_text}</div>
                        </div>
                        """, unsafe_allow_html=True)
        
        with col3:
            with st.container(border=True):
                st.markdown("#### 📊 Agreement ")
                
                # Calculate summary metrics  
                total_methods = len(results)
                flagged_count = sum(flags.values())
                passed_count = total_methods - flagged_count
                
                summary_cols = st.columns(3)
                
                with summary_cols[0]:
                    st.metric("Methods", total_methods)
                
                with summary_cols[1]:
                    st.metric("Flagged", flagged_count,
                            delta=f"-{flagged_count}" if flagged_count > 0 else "0",
                            delta_color="inverse")
                
                with summary_cols[2]:
                    st.metric("Passed", passed_count,
                            delta=f"+{passed_count}" if passed_count > 0 else "0",
                            delta_color="normal")
                st.markdown(f"""
                <div class="text-align" style="background-color: {agreement_color}20; border-left: 0px solid {agreement_color};">
                    <strong>Agreement Confidence: {agreement_confidence}</strong><br>
                    {agreement_level} ({agreement_ratio})<br>
                </div>
                """, unsafe_allow_html=True)
        
        # Score Improvement Guidance
        # st.markdown('<div class="section-header">💡 Score Improvement Guidance</div>', unsafe_allow_html=True)
        
        # Calculate guidance based on current status
        guidance_message = ""
        guidance_color = ""
        
        # Get score range for current test
        min_score = score_ranges.get(current_score, {}).get('min', 0)
        max_score = score_ranges.get(current_score, {}).get('max', 100)
        
        if best_flagged:
            # User is below threshold - find score needed to cross threshold
            test_score = int(current_actual) + 1
            found_threshold_score = None
            last_valid_percentile = best_percentile  # Track last valid percentile for monotonicity
            max_achievable_percentile = best_percentile  # Track the maximum achievable percentile
            
            # First, find the score that crosses the threshold
            while test_score <= max_score:
                # Calculate percentile for this test score with monotonicity constraint
                test_results, _, test_best_method, test_best_percentile, _ = model_manager.calculate_percentile_with_best_method(
                    current_score, float(test_score), features, min_percentile_constraint=last_valid_percentile
                )
                
                if test_best_percentile:
                    max_achievable_percentile = max(max_achievable_percentile, test_best_percentile)
                    
                    if test_best_percentile >= current_threshold and found_threshold_score is None:
                        found_threshold_score = test_score
                        found_percentile = test_best_percentile
                
                # Update last valid percentile if we found a higher one
                if test_best_percentile and test_best_percentile > last_valid_percentile:
                    last_valid_percentile = test_best_percentile
                
                test_score += 1
            
            # Also check what percentile can be achieved at maximum score
            # For maximum score, we want the highest possible percentile, not the best coverage
            max_score_results, _, max_score_method, max_score_percentile, _ = model_manager.calculate_percentile_with_best_method(
                current_score, float(max_score), features, 
                min_percentile_constraint=last_valid_percentile,
                use_max_percentile=True  # Select method with highest percentile for max score
            )
            
            if max_score_percentile:
                max_achievable_percentile = max(max_achievable_percentile, max_score_percentile)
            
            if found_threshold_score:
                guidance_message = f"""
                <div class="info-box-alert">
                    <strong>🎯 To Cross the Threshold:</strong><br>
                    You need to score at least <strong>{found_threshold_score}</strong> to exceed the {current_threshold}th percentile threshold.<br>
                    <small>A score of {found_threshold_score} would place you at approximately the {found_percentile:.0f}th percentile.</small>
                </div>
                """
            elif max_achievable_percentile >= current_threshold:
                # Threshold can be reached, but we need to find the exact score
                # This handles cases where sequential checking missed it due to monotonicity constraints
                guidance_message = f"""
                <div class="info-box-alert">
                    <strong>🎯 Score Guidance:</strong><br>
                    The {current_threshold}th percentile threshold can be achieved with a higher score.<br>
                    <small>Maximum achievable percentile: {max_achievable_percentile:.0f}th at score {max_score} (using {max_score_method})</small>
                </div>
                """
            else:
                guidance_message = f"""
                <div class="info-box-alert">
                    <strong>ℹ️ Score Guidance:</strong><br>
                    Based on your demographic profile, the maximum achievable percentile is <strong>{max_achievable_percentile:.0f}th</strong> 
                    (at score {max_score}, using {max_score_method}), which is below the {current_threshold}th percentile threshold.
                </div>
                """
        else:
            # User is above threshold - show next score improvement
            next_score = int(current_actual) + 1
            
            if next_score <= max_score:
                # Calculate percentile for next score with monotonicity constraint
                next_results, _, next_best_method, next_best_percentile, _ = model_manager.calculate_percentile_with_best_method(
                    current_score, float(next_score), features, min_percentile_constraint=best_percentile
                )
                
                if next_best_percentile:
                    improvement = next_best_percentile - best_percentile
                    guidance_message = f"""
                    <div class="info-box-success">
                        <strong>🚀 Score Improvement Potential:</strong><br>
                        If you score <strong>{next_score}</strong> (one point higher), you would be at approximately the <strong>{next_best_percentile:.0f}th percentile</strong>.<br>
                        <small>This represents a {improvement:.0f} percentile point improvement from your current position.</small>
                    </div>
                    """
            else:
                guidance_message = f"""
                <div class="info-box-success">
                    <strong>⭐ Maximum Score Achieved:</strong><br>
                    You have achieved the maximum possible raw score ({max_score}) for this test.
                </div>
                """
        
        # Display guidance message
        if guidance_message:
            st.markdown(guidance_message, unsafe_allow_html=True)
        
        # Action buttons
        # st.markdown("---")
        col1, col2, col3 = st.columns([1, 1, 3])
        
        with col1:
            # Save to session
            if st.button("💾 Save Test Result", type="primary"):
                # Create a unique identifier for this specific test configuration
                test_config_id = f"{current_score}_{current_actual}_{current_threshold}_{best_method}_{best_percentile}"
                
                # Check if this exact configuration was already saved
                already_saved = any(
                    f"{test['test_name']}_{test['actual_score']}_{test['threshold']}_{test['best_method']}_{test['best_percentile']}" == test_config_id
                    for test in st.session_state.all_test_results
                )
                
                if not already_saved:
                    # Compile all test data
                    test_result = {
                        'test_name': current_score,
                        'actual_score': current_actual,
                        'threshold': current_threshold,
                        'best_method': best_method,
                        'best_percentile': best_percentile,
                        'flag_status': 'FLAGGED' if best_flagged else 'PASS',
                        # 'interpretation': interpretation,
                        'agreement_level': agreement_level,
                        'agreement_ratio': agreement_ratio,
                        'agreement_confidence': agreement_confidence,
                        'all_methods': {
                            method: {
                                'percentile': results[method],
                                'flag': 'FLAGGED' if flags[method] else 'PASS',
                                'offset': offsets.get(method, 0)
                            }
                            for method in results
                        },
                        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    }
                    
                    # Add to session
                    st.session_state.all_test_results.append(test_result)
                    st.session_state.test_counter += 1
                    st.success(f"✅ Test result saved! ({len(st.session_state.all_test_results)} tests in session)")
                    # Force rerun to update sidebar immediately
                    st.rerun()
                else:
                    st.warning("⚠️ This exact test configuration has already been saved! Change the threshold or other parameters to save a different configuration.")
        
        with col2:
            if st.button("🔄 New Test"):
                # Clear current results
                if 'current_results' in st.session_state:
                    del st.session_state.current_results
                st.rerun()
            
    elif calculate_btn:
        st.warning("⚠️ Please enter an actual score value.")
    
    else:
        # Instructions when no results
        st.markdown('<div class="section-header">📋 Quick Start Guide</div>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            <div class="info-box">
            <strong>1. Demographics</strong><br>
            Enter subject details in the sidebar
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="info-box">
            <strong>2. Test Settings</strong><br>
            Select test and set threshold above
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="info-box">
            <strong>3. Calculate</strong><br>
            Enter score and click Calculate
            </div>
            """, unsafe_allow_html=True)
    
    # Disclaimer - Always visible at the bottom
    st.markdown("---")
    st.markdown("""
    <div style="background-color: #fff3cd; border: 1px solid #ffeaa7; border-radius: 8px; 
                padding: 1rem; margin-top: -0.2rem; margin-bottom: 2rem; color: #856404;">
        <strong>⚠️ Important Disclaimer:</strong><br>
        This application is intended primarily as a research tool for exploring and comparing normative modeling approaches. 
        Users are advised to interpret results with caution and avoid relying solely on its outputs for critical clinical decisions. 
        Final interpretation should always be guided by clinical expertise, corroborative assessments, and context-specific considerations.
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()